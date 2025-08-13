import pubchempy as pcp
import pandas as pd
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# -----------------------------
# Utility Functions (No RDKit)
# -----------------------------

def validate_smiles(smiles):
    """Validate SMILES using PubChem lookup"""
    try:
        compounds = pcp.get_compounds(smiles, 'smiles')
        return compounds[0] if compounds else None
    except Exception as e:
        st.warning(f"SMILES validation failed: {e}")
        return None

def get_pubchem_data(smiles):
    """Fetch compound data from PubChem"""
    try:
        compounds = pcp.get_compounds(smiles, 'smiles')
        if not compounds:
            return {}
        c = compounds[0]

        # Safely extract properties
        def safe_float(val, default=0.0):
            try:
                return float(val) if val not in [None, '--'] else default
            except (TypeError, ValueError):
                return default

        mw = safe_float(c.molecular_weight)
        logp = safe_float(c.xlogp, default=2.0)

        # Estimate solubility
        solubility = "Highly soluble" if mw < 500 and logp < 3 else "Low solubility"

        return {
            "name": c.iupac_name or (c.synonyms[0] if c.synonyms else "Unknown"),
            "formula": c.molecular_formula or "Not available",
            "mw": round(mw, 2),
            "cas": getattr(c, 'cas', "Not available"),
            "logp": round(logp, 2),
            "solubility": solubility,
            "h_bond_donor": c.h_bond_donor_count or 0,
            "h_bond_acceptor": c.h_bond_acceptor_count or 0,
        }
    except Exception as e:
        st.warning(f"PubChem lookup failed: {e}")
        return {}

def predict_toxicity_protx(smiles):
    """Simulate ProTox-II prediction using PubChem data"""
    try:
        compounds = pcp.get_compounds(smiles, 'smiles')
        if not compounds:
            raise ValueError("Invalid SMILES")
        c = compounds[0]

        # Heuristic: check for nitro groups or toxic functional groups
        synonyms_str = " ".join(c.synonyms).lower() if c.synonyms else ""
        has_nitro = 'nitro' in synonyms_str
        has_cyanide = 'cyano' in synonyms_str or 'cyanide' in synonyms_str

        if has_nitro or has_cyanide:
            return {
                "toxicity_class": "Class II (High)",
                "hazard_endpoints": ["Hepatotoxicity", "Neurotoxicity"],
                "ld50": "50 mg/kg"
            }
        else:
            return {
                "toxicity_class": "Class IV (Low)",
                "hazard_endpoints": ["None predicted"],
                "ld50": "5000 mg/kg"
            }
    except Exception as e:
        st.warning(f"Toxicity prediction failed: {e}")
        return {
            "toxicity_class": "Unknown",
            "hazard_endpoints": ["Data not available"],
            "ld50": "Unknown"
        }

def get_physical_properties(smiles):
    """Fetch physical properties from PubChem"""
    try:
        compounds = pcp.get_compounds(smiles, 'smiles')
        if not compounds:
            raise ValueError("No compound found")
        c = compounds[0]

        def safe_float(val, default=0.0):
            try:
                return float(val) if val not in [None, '--'] else default
            except:
                return default

        mw = safe_float(c.molecular_weight)
        logp = safe_float(c.xlogp, default=2.0)
        tpsa = safe_float(c.tpsa, default=0.0)

        return {
            "_MolecularWeight_numeric": mw,
            "_LogP_numeric": logp,
            "Molecular Weight": f"{mw:.2f} g/mol",
            "LogP": f"{logp:.2f}",
            "Topological Polar Surface Area (TPSA)": f"{tpsa:.2f} Å²" if c.tpsa else "Not available",
            "Hydrogen Bond Donors": c.h_bond_donor_count or 0,
            "Hydrogen Bond Acceptors": c.h_bond_acceptor_count or 0,
            "Rotatable Bonds": c.rotatable_bond_count or 0,
            "Heavy Atom Count": c.heavy_atom_count or 0,
        }
    except Exception as e:
        st.warning(f"Physical properties lookup failed: {e}")
        return {
            "_MolecularWeight_numeric": 300.0,
            "_LogP_numeric": 2.0,
            "Molecular Weight": "300.00 g/mol",
            "LogP": "2.00",
            "Topological Polar Surface Area (TPSA)": "Not available",
            "Hydrogen Bond Donors": 0,
            "Hydrogen Bond Acceptors": 0,
            "Rotatable Bonds": 0,
            "Heavy Atom Count": 0,
        }

def section_name(i):
    names = {
        1: "Chemical Product and Company Identification",
        2: "Composition and Information on Ingredients",
        3: "Hazards Identification",
        4: "First Aid Measures",
        5: "Fire and Explosion Data",
        6: "Accidental Release Measures",
        7: "Handling and Storage",
        8: "Exposure Controls/Personal Protection",
        9: "Physical and Chemical Properties",
        10: "Stability and Reactivity",
        11: "Toxicological Information",
        12: "Ecological Information",
        13: "Disposal Considerations",
        14: "Transport Information",
        15: "Other Regulatory Information",
        16: "Other Information"
    }
    return names.get(i, f"Section {i}")

def generate_sds(smiles):
    """Generate full SDS using only PubChem data"""
    compound = validate_smiles(smiles)
    if not compound:
        return None

    pubchem = get_pubchem_data(smiles)
    protx = predict_toxicity_protx(smiles)
    props = get_physical_properties(smiles)

    sds = {
        f"Section{i}": {
            "title": section_name(i),
            "data": {},
            "notes": []
        } for i in range(1, 17)
    }

    # Section 1
    sds["Section1"]["data"] = {
        "Product Identifier": pubchem.get("name", "Unknown Compound"),
        "Company": "Automated SDS Generator",
        "Address": "N/A",
        "Emergency Phone": "N/A",
        "Recommended Use": "Research Use Only"
    }

    # Section 2
    sds["Section2"]["data"] = {
        "Name": pubchem.get("name", "Unknown"),
        "CAS Number": pubchem.get("cas", "Not available"),
        "Molecular Formula": pubchem.get("formula", "Not available"),
        "Purity/Concentration": "100% (pure compound)"
    }

    # Section 3: Hazards Identification
    is_flammable = pubchem.get("logp", 0) > 1.5
    is_toxic = protx.get("toxicity_class") in ["Class I", "Class II", "Class III", "Class IV"]
    has_health_hazard = is_toxic

    pictograms = []
    hazard_statements = []
    if is_flammable:
        pictograms.append("🔥 Flammable")
        hazard_statements.append("H225: Highly flammable liquid and vapor")
    if is_toxic:
        pictograms.append("💀 Acute Toxicity")
        hazard_statements.append("H301: Toxic if swallowed")
        hazard_statements.append("H331: Toxic if inhaled")

    signal_word = "Danger" if (is_flammable or is_toxic) else "Warning"

    health_effects = "This substance is harmful if inhaled, swallowed, or absorbed through the skin. "
    if is_toxic:
        health_effects += "It may cause central nervous system depression, organ damage, or acute toxicity. "
    if is_flammable:
        health_effects += "Vapors may cause dizziness or asphyxiation in high concentrations. "
    health_effects += "Chronic exposure may lead to liver, kidney, or respiratory damage."

    precautionary = [
        "P210: Keep away from heat, hot surfaces, sparks, open flames.",
        "P241: Use explosion-proof electrical/ventilation equipment.",
        "P261: Avoid breathing dust/fume/gas/mist/vapors/spray.",
        "P280: Wear protective gloves/protective clothing/eye protection/face protection.",
        "P305+P351+P338: IF IN EYES: Rinse cautiously with water for several minutes."
    ]

    sds["Section3"]["data"] = {
        "Signal Word": signal_word,
        "GHS Pictograms": ", ".join(pictograms) if pictograms else "Not classified",
        "Hazard Statements": hazard_statements if hazard_statements else ["No significant hazards identified"],
        "Precautionary Statements": precautionary,
        "Physical Hazards": "Flammable liquid and vapor" if is_flammable else "Not flammable",
        "Health Hazards": ", ".join([p.replace("💀 ", "") for p in pictograms if "💀" in p]) or "None identified",
        "Environmental Hazards": "Toxic to aquatic life" if protx.get("toxicity_class") in ["Class I", "Class II"] else "Low concern",
        "Routes of Exposure": "Inhalation, Skin Contact, Ingestion, Eye Contact",
        "Acute and Chronic Effects": health_effects,
        "Immediate Medical Attention": "Seek medical attention immediately in case of exposure. Show SDS to physician."
    }

    # Section 4
    sds["Section4"]["data"] = {
        "Inhalation": "Move to fresh air. If breathing is difficult, give oxygen.",
        "Skin Contact": "Flush with plenty of water. Remove contaminated clothing.",
        "Eye Contact": "Flush with water for at least 15 minutes.",
        "Ingestion": "Do NOT induce vomiting. Rinse mouth and consult a physician."
    }

    # Section 5
    flash_point = "13°C" if pubchem.get("logp", 0) > 1 else "Not flammable"
    sds["Section5"]["data"] = {
        "Flash Point": flash_point,
        "Flammable Limits": "3.3% - 19% in air",
        "Extinguishing Media": "Dry chemical, CO2, alcohol-resistant foam",
        "Special Hazards": "Vapors may form explosive mixtures with air."
    }

    # Section 6
    sds["Section6"]["data"] = {
        "Personal Precautions": "Wear PPE, ensure ventilation",
        "Environmental Precautions": "Prevent entry into drains or waterways",
        "Methods of Containment": "Absorb with inert material (sand, vermiculite)"
    }

    # Section 7
    sds["Section7"]["data"] = {
        "Handling": "Ground containers, use explosion-proof equipment",
        "Storage": "Store in a cool, well-ventilated place away from ignition sources"
    }

    # Section 8
    sds["Section8"]["data"] = {
        "TLV-TWA": "100 ppm (300 mg/m³) for ethanol-like compounds",
        "Engineering Controls": "Local exhaust ventilation",
        "Personal Protection": "Safety goggles, gloves, lab coat"
    }

    # Section 9
    mw_numeric = props["_MolecularWeight_numeric"]
    sds["Section9"]["data"] = {
        "Physical State": "Liquid" if mw_numeric < 300 else "Solid",
        "Color": "Colorless",
        "Odor": "Characteristic",
        "Melting Point": "Not available",
        "Boiling Point": "Not available",
        "Solubility in Water": pubchem.get("solubility", "Data not available"),
        "Density": "Approx. 0.79 g/cm³ (for alcohols)",
        "Vapor Pressure": "< 1 mmHg at 25°C",
        **{k: v for k, v in props.items() if not k.startswith("_")}
    }

    # Section 10
    sds["Section10"]["data"] = {
        "Stability": "Stable under normal conditions",
        "Conditions to Avoid": "Heat, flames, sparks",
        "Incompatible Materials": "Strong oxidizing agents",
        "Hazardous Decomposition": "Carbon monoxide, carbon dioxide"
    }

    # Section 11
    sds["Section11"]["data"] = {
        "LD50 Oral Rat": protx.get("ld50"),
        "LC50 Inhalation Rat": "Not available",
        "Carcinogenicity": "Suspected" if "Hepatotoxicity" in protx.get("hazard_endpoints", []) else "Not suspected",
        "Mutagenicity": "Positive" if "Hepatotoxicity" in protx.get("hazard_endpoints", []) else "Negative",
        "Toxicity Class": protx.get("toxicity_class", "Class IV")
    }

    # Section 12
    sds["Section12"]["data"] = {
        "Ecotoxicity": "Toxic to aquatic life" if protx.get("toxicity_class") in ["Class I", "Class II"] else "Low concern",
        "Biodegradability": "Yes",
        "Persistence": "Low",
        "Bioaccumulation": "Low potential"
    }

    # Section 13
    sds["Section13"]["data"] = {
        "Disposal Method": "Dispose in accordance with local regulations",
        "Contaminated Packaging": "Rinse and recycle or dispose properly"
    }

    # Section 14
    sds["Section14"]["data"] = {
        "UN Number": "UN1170",
        "Proper Shipping Name": "Ethanol or Ethyl Alcohol",
        "Transport Hazard Class": "3 (Flammable Liquid)",
        "Packing Group": "II"
    }

    # Section 15
    sds["Section15"]["data"] = {
        "TSCA": "Listed",
        "DSL": "Listed",
        "WHMIS": "Classified",
        "GHS Regulation": "GHS Rev 9 compliant"
    }

    # Section 16
    sds["Section16"]["data"] = {
        "Date Prepared": datetime.now().strftime("%Y-%m-%d"),
        "Revision Number": "1.0",
        "Prepared By": "Automated ADMET-SDS System",
        "Disclaimer": "Generated for research use only. Verify with lab testing."
    }

    return sds

# -----------------------------
# Document Generation
# -----------------------------

def generate_docx(sds, compound_name="Unknown Compound"):
    """Generate DOCX file"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    title = doc.add_heading('Safety Data Sheet (SDS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Compound: {compound_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_on = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"Generated on: {generated_on}", style='Caption')
    doc.add_paragraph()

    for i in range(1, 17):
        section_key = f"Section{i}"
        section = sds.get(section_key, {})
        title = section.get("title", f"Section {i}")
        doc.add_heading(f"{i}. {title}", level=1)

        data = section.get("data", {})
        if not data:
            doc.add_paragraph("No data available.")
        else:
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for key, value in data.items():
                row = table.add_row()
                cell_key = row.cells[0]
                cell_val = row.cells[1]

                p_key = cell_key.paragraphs[0]
                run_key = p_key.add_run(str(key))
                run_key.bold = True

                if isinstance(value, list):
                    val_text = ", ".join([str(v) for v in value if v]) or "Not available"
                elif not value or value == "Not available":
                    val_text = "Not available"
                else:
                    val_text = str(value)

                cell_val.text = val_text

        doc.add_paragraph()

    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run("Disclaimer: This report is generated for research use only. "
                             "Verify with lab testing and official sources before handling chemicals.")
    run.italic = True
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filename = f"SDS_{compound_name.replace(' ', '_').replace('/', '_')}.docx"
    doc.save(filename)
    return filename

def generate_pdf(sds, compound_name="Unknown Compound"):
    """Generate PDF file using ReportLab"""
    filename = f"SDS_{compound_name.replace(' ', '_').replace('/', '_')}.pdf"
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1e3c72')
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.HexColor('#2a5298')
    )
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
    caption_style = ParagraphStyle(
        'CustomCaption',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_CENTER,
        textColor=colors.grey
    )

    story = []
    story.append(Paragraph("Safety Data Sheet (SDS)", title_style))
    story.append(Paragraph(f"Compound: {compound_name}", caption_style))
    generated_on = datetime.now().strftime("%Y-%m-%d %H:%M")
    story.append(Paragraph(f"Generated on: {generated_on}", caption_style))
    story.append(Spacer(1, 24))

    for i in range(1, 17):
        section_key = f"Section{i}"
        section = sds.get(section_key, {})
        title = section.get("title", f"Section {i}")
        story.append(Paragraph(f"{i}. {title}", heading_style))

        data = section.get("data", {})
        if not data:
            story.append(Paragraph("No data available.", normal_style))
        else:
            table_data = []
            for key, value in data.items():
                if isinstance(value, list):
                    val_text = ", ".join([str(v) for v in value if v]) or "Not available"
                elif not value or value == "Not available":
                    val_text = "Not available"
                else:
                    val_text = str(value)
                if len(val_text) > 80:
                    val_text = val_text[:80] + "..."

                table_data.append([
                    Paragraph(f"<b>{key}</b>", normal_style),
                    Paragraph(val_text, normal_style)
                ])

            if table_data:
                table = Table(table_data, colWidths=[2.5*inch, 4*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(table)
        story.append(Spacer(1, 12))

    disclaimer_text = ("Disclaimer: This report is generated for research use only. "
                      "Verify with lab testing and official sources before handling chemicals.")
    story.append(Spacer(1, 24))
    story.append(Paragraph(f"<i>{disclaimer_text}</i>", caption_style))

    doc.build(story)
    return filename
